import os
import io
import re
import uuid
from typing import List
from fastapi import FastAPI, HTTPException, Depends, status, File, UploadFile, Form
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from pydantic import BaseModel
from google import genai
from google.genai import types
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
import httpx

load_dotenv()
app = FastAPI()
security = HTTPBasic()

# OpenClawのローカルGateway用URL
OPENCLAW_API_URL = os.getenv("OPENCLAW_API_URL", "http://localhost:18789/api/v1")

# 📸 画像を一時公開・保持するためのディレクトリ作成とマウント設定
UPLOAD_DIR = "uploaded_images"
os.makedirs(UPLOAD_DIR, exist_ok=True)
app.mount("/images", StaticFiles(directory=UPLOAD_DIR), name="images")

def authenticate(credentials: HTTPBasicCredentials = Depends(security)):
    correct_username = os.getenv("WEB_USERNAME", "1")
    correct_password = os.getenv("WEB_PASSWORD", "1")
    if credentials.username != correct_username or credentials.password != correct_password:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="ユーザー名またはパスワードが違います",
            headers={"WWW-Authenticate": "Basic"},
        )
    return credentials.username

try:
    client = genai.Client()
except Exception as e:
    print(f"Geminiの初期化に失敗しました: {e}")

class DownloadRequest(BaseModel):
    title: str
    result: str

# ----------------------------------------------------------------
# 🦞 OpenClaw 自律クローリング実行ヘルパー
# ----------------------------------------------------------------
async def run_openclaw_research(title: str) -> str:
    prompt_for_claw = f"「{title}」について、電気工学・実験原理の観点から最新の技術動向、回路構成、数式的な背景をWebブラウザ等で自動検索・クローリングし, 詳細な調査結果を日本語のMarkdownテキストとしてまとめてください。"
    
    async with httpx.AsyncClient(timeout=120.0) as http_client:
        try:
            response = await http_client.post(
                f"{OPENCLAW_API_URL}/agent/run",
                json={
                    "message": prompt_for_claw,
                    "skills": ["web_search", "browser_fetch"]
                }
            )
            if response.status_code == 200:
                data = response.json()
                return data.get("output", "OpenClawからの応答が空でした。")
            else:
                return f"【警告】OpenClaw連携エラー (Status: {response.status_code})"
        except Exception as e:
            return f"【システムノート】OpenClawデーモンが不通のため、標準検索でフォールバックします。 (詳細: {str(e)})"

# ----------------------------------------------------------------
# 📐 Word用の数式パース＆書き込みヘルパー
# ----------------------------------------------------------------
def add_math_to_paragraph(paragraph, latex_str):
    try:
        clean_latex = latex_str.strip().replace("\n", "")
        if clean_latex:
            run = paragraph.add_run(f" {clean_latex} ")
            run.font.name = 'Cambria Math'
            run.italic = True
    except:
        paragraph.add_run(f" {latex_str} ")

def parse_and_write_text_with_math(paragraph, text):
    """テキスト中の $...$ または $$...$$（数式）をパースして適切にRunを追加する"""
    parts = re.split(r'(\$\$.*?\$\$|\$.*?\$)', text, flags=re.DOTALL)
    for part in parts:
        if part.startswith('$$') and part.endswith('$$'):
            latex_formula = part[2:-2]
            add_math_to_paragraph(paragraph, latex_formula)
        elif part.startswith('$') and part.endswith('$'):
            latex_formula = part[1:-1]
            add_math_to_paragraph(paragraph, latex_formula)
        else:
            if part:
                paragraph.add_run(part)

# ----------------------------------------------------------------
# 🌐 各種APIエンドポイント
# ----------------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
async def read_index():
    with open("index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.get("/auth-check")
async def auth_check(username: str = Depends(authenticate)):
    return {"status": "success", "username": username}

@app.post("/process-report")
async def process_report(
    title: str = Form(...),
    report_type: str = Form(...),
    instruction: str = Form(""),
    files: List[UploadFile] = File(...),
    username: str = Depends(authenticate)
):
    ALLOWED_MIMETYPES = ["application/pdf", "image/jpeg", "image/png", "image/webp"]
    
    gemini_files = []
    saved_image_mappings = []
    
    try:
        for file in files:
            if file.filename == "": continue
            file_bytes = await file.read()
            
            if file.content_type in ["image/jpeg", "image/png", "image/webp"]:
                ext = os.path.splitext(file.filename)[1]
                unique_filename = f"{uuid.uuid4()}{ext}"
                file_path = os.path.join(UPLOAD_DIR, unique_filename)
                with open(file_path, "wb") as f:
                    f.write(file_bytes)
                
                saved_image_mappings.append({
                    "original": file.filename,
                    "unique": unique_filename
                })
            
            gemini_file = client.files.upload(
                file=io.BytesIO(file_bytes),
                config=types.UploadFileConfig(mime_type=file.content_type)
            )
            gemini_files.append(gemini_file)
        
        crawled_knowledge = await run_openclaw_research(title)

        if report_type == "experiment":
            type_prompt = """
            【レポート構成および記述内容の厳格な指定】
            以下の構成に沿って、各章とも学術論文レベルの圧倒的な情報量と厳密な専門用語（〜である調）で詳細に執筆してください。
            章タイトル（大見出し）の行頭には必ず [H1] を、節タイトル（中見出し）の行頭には必ず [H2] を付与してください。

            [H1] 1. 実験目的
            [H1] 2. 実験原理
            ・[H2] 2.1 フーリエ級数展開
            ・[H2] 2.2 複素フーリエ級数
            ・[H2] 2.3 フーリエ変換・逆フーリエ変換
            ・[H2] 2.4 DFTとFFT
            ・[H2] 2.5 フーリエ解析の性質
            ・[H2] 2.6 リーケージ現象と窓関数
            [H1] 3. 実験方法および回路構成
            ・[H2] 3.1 多成分正弦波の合成
            ・[H2] 3.2 FFT処理と片側振幅スペクトルの算出
            ・[H2] 3.3 窓関数実験手順
            ・[H2] 3.4 非正弦周期信号の解析手順
            [H1] 4. 実験結果
            ・[H2] 4.1 合成信号の解析結果
            ・[H2] 4.2 窓関数実験の結果
            ・[H2] 4.3 非正弦波形の結果
            [H1] 5. 考察
            ・[H2] 5.1 周波数分離特性の限界
            ・[H2] 5.2 リーケージと窓関数の数理的考察
            ・[H2] 5.3 高調波発生と波形対称性・連続性の考察
            ・[H2] 5.4 サンプリング定理とエイリアシング
            [H1] 6. 結論
            """
        elif report_type == "application_advanced":
            type_prompt = """
            【レポート構成および記述内容の厳格な指定】
            [H1] 1. 総合概要 (Executive Summary)
            [H1] 2. 技術的背景および基本原理
            [H1] 3. 発発展回路（応用回路）の構成と動作解析
            [H1] 4. 現在の主な応用例・社会実装動向
            [H1] 5. 現状の技術的課題とその対策
            [H1] 6. 今後の展望および提言
            """
        else:
            type_prompt = """
            【レポート構成および記述内容の厳格な指定】
            [H1] 1. 総合概要 (Executive Summary)
            [H1] 2. 技術的背景および基本原理
            [H1] 3. 現在の主な応用例・社会実装動向
            [H1] 4. 現状の技術的課題とその対策
            [H1] 5. 今後の展望および提言
            """

        image_instruction_str = ""
        if saved_image_mappings:
            image_instruction_str = "【利用可能な画像資料の一覧】:\n"
            for mapping in saved_image_mappings:
                image_instruction_str += f"・ユーザーが添付した写真「{mapping['original']}」 ➡️ 挿入用指定タグ: [IMAGE: {mapping['unique']}]\n"

        prompt = f"""
        あなたはプロフェッショナルな電気工学アナリスト兼リサーチライターです。
        提供された添付資料に加え、自律エージェント（OpenClaw）がWebから自動クローリングしてきた技術データ、さらにユーザーが添付した写真データを完全に融合させて高品質なレポートを執筆してください。
        
        {image_instruction_str}
        
        【作成するレポートのタイトル】: {title}
        【追加の作成指示・メモ】: {instruction}
        {type_prompt}
        
        【超重要ルール：画像の自動挿入配置】
        ・上記に「利用可能な画像資料の一覧」がある場合、その内容を分析し、レポートの文脈に合わせて最も適切な位置に、該当する画像の挿入用指定タグ（例: `[IMAGE: xxx-xxx.ext]`) を【必ず単独の行】として配置してください。
        
        【超重要ルール：数式のアカデミック化】
        ・数式や理論式を出力する場合は、必ずインライン数式（$E=mc^2$）または独立数式（$$数式$$）のLaTeX表記を使用してください。
        
        【超重要ルール：データ羅列の表形式化】
        ・測定データや解析結果の箇条書きの羅列を出力にそのまま含めず、必ず人間が見やすい「Markdownの表（テーブル）形式」へと自動的に変換・整理して出力してください。
        
        ・AIのメタ発言、前置き、結びの挨拶は【絶対に】含めず、純粋なレポートの章と本文だけを出力してください。
        """
        
        contents = gemini_files + [prompt]
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=contents,
            config=types.GenerateContentConfig(tools=[{"google_search": {}}])
        )
        
        for g_file in gemini_files:
            try: client.files.delete(name=g_file.name)
            except: pass
        
        return {"status": "success", "result": response.text}
        
    except Exception as e:
        for g_file in gemini_files:
            try: client.files.delete(name=g_file.name)
            except: pass
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/edit-report")
async def edit_report(
    title: str = Form(...),
    current_result: str = Form(...),
    edit_instruction: str = Form(...),
    files: List[UploadFile] = File(default=[]),
    username: str = Depends(authenticate)
):
    ALLOWED_MIMETYPES = ["application/pdf", "image/jpeg", "image/png", "image/webp"]
    gemini_files = []
    saved_image_mappings = []
    try:
        for file in files:
            if file.filename == "": continue
            file_bytes = await file.read()
            
            if file.content_type in ["image/jpeg", "image/png", "image/webp"]:
                ext = os.path.splitext(file.filename)[1]
                unique_filename = f"{uuid.uuid4()}{ext}"
                file_path = os.path.join(UPLOAD_DIR, unique_filename)
                with open(file_path, "wb") as f:
                    f.write(file_bytes)
                saved_image_mappings.append({
                    "original": file.filename,
                    "unique": unique_filename
                })
            
            gemini_file = client.files.upload(
                file=io.BytesIO(file_bytes),
                config=types.UploadFileConfig(mime_type=file.content_type)
            )
            gemini_files.append(gemini_file)

        image_instruction_str = ""
        if saved_image_mappings:
            image_instruction_str = "【新しく追加添付された画像資料の一覧】:\n"
            for mapping in saved_image_mappings:
                image_instruction_str += f"・追加された写真「{mapping['original']}」 ➡️ 挿入用指定タグ: [IMAGE: {mapping['unique']}]\n"

        prompt = f"""
        あなたは提出前のレポートを極限までブラッシュアップする優秀な校正・編集者です。
        現在作成されている以下の【現在のレポート】に対して、ユーザーから【修正・追加指示】および【新しく追加添付された参考資料ファイル】が届きました。
        
        {image_instruction_str}
        
        現在の構成や本文をベースにしつつ、新データや指示された内容を完全に反映した「修正後の最新レポート」を再出力してください。
        
        【レポートタイトル】: {title}
        【現在のレポート】:
        {current_result}
        【ユーザーからの修正・追加指示】:
        {edit_instruction}
        
        【超重要ルール】
        ・新しく画像が追加された場合は、適切な位置に `[IMAGE: xxx.ext]` の形式で【単独行】として必ず画像を配置してください。
        ・章タイトルの行頭には必ず [H1] を、節タイトルの行頭には必ず [H2] を維持または付与してください。
        ・数式は必ず1行の $...$ または $$...$$ のLaTeX形式を死守してください。
        ・出力はそのままWordに変換されます。案内文や前置きは【絶対に】入れないでください。
        """
        
        contents = gemini_files + [prompt]
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=contents,
            config=types.GenerateContentConfig(tools=[{"google_search": {}}])
        )
        
        for g_file in gemini_files:
            try: client.files.delete(name=g_file.name)
            except: pass
            
        return {"status": "success", "result": response.text}
    except Exception as e:
        for g_file in gemini_files:
            try: client.files.delete(name=g_file.name)
            except: pass
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/download/docx")
async def download_docx(request: DownloadRequest, username: str = Depends(authenticate)):
    doc = Document()
    
    # 📄 【新設】1枚目を独立した「表紙専用ページ」にする自動改ページロジック
    for _ in range(6): # タイトルの上に適度な余白（空行）を作って中央に寄せる
        doc.add_paragraph()
        
    title_heading = doc.add_heading(request.title, level=1)
    title_heading.alignment = 1 # タイトル文字を中央揃えに配置
    
    doc.add_page_break() # 👈 ここで強制的にページを区切ることで、1枚目を完璧な表紙にする！
    
    lines = request.result.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        trimmed = line.strip()
        
        if not trimmed:
            i += 1
            continue
            
        if trimmed.startswith("[H1] "):
            doc.add_heading(trimmed.replace("[H1] ", ""), level=2)
            i += 1
        elif trimmed.startswith("[H2] "):
            doc.add_heading(trimmed.replace("[H2] ", ""), level=3)
            i += 1
        elif trimmed.startswith("[IMAGE: ") and trimmed.endswith("]"):
            img_filename = trimmed.replace("[IMAGE: ", "").replace("]", "")
            img_path = os.path.join(UPLOAD_DIR, img_filename)
            if os.path.exists(img_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = 1
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(4.5))
                except:
                    doc.add_paragraph(f"【画像の展開失敗: {img_filename}】")
            else:
                doc.add_paragraph(f"【画像ファイル紛失: {img_filename}】")
            i += 1
        elif trimmed.startswith("- ") or trimmed.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            parse_and_write_text_with_math(p, line[2:])
            i += 1
        elif trimmed.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                t_line = lines[i].strip()
                if not re.match(r'^\|[\s*-:|]*\|$', t_line):
                    table_lines.append(t_line)
                i += 1
            
            if table_lines:
                matrix = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.strip('|').split('|')]
                    matrix.append(cells)
                
                if matrix:
                    cols_num = max(len(row) for row in matrix)
                    table = doc.add_table(rows=len(matrix), cols=cols_num)
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(matrix):
                        word_row = table.rows[row_idx]
                        for col_idx, cell_value in enumerate(row_data):
                            if col_idx < len(word_row.cells):
                                p = word_row.cells[col_idx].paragraphs[0]
                                parse_and_write_text_with_math(p, cell_value)
        else:
            p = doc.add_paragraph()
            parse_and_write_text_with_math(p, line)
            i += 1
                
    file_io = io.BytesIO()
    doc.save(file_io)
    file_io.seek(0)
    return StreamingResponse(
        file_io,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=report.docx"}
    )